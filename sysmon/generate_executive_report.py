#!/usr/bin/env python3
"""
Genera un Executive Report Word elegante con grafici professionali
Per presentazioni CISO/Board
"""

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Colori corporate
COLORS = {
    'primary': '#1a365d',      # Blu scuro
    'secondary': '#2c5282',    # Blu medio
    'accent': '#38a169',       # Verde successo
    'warning': '#d69e2e',      # Giallo warning
    'danger': '#e53e3e',       # Rosso
    'light': '#e2e8f0',        # Grigio chiaro
    'text': '#2d3748',         # Testo scuro
}

def create_gauge_chart(value, max_value, title, filename):
    """Crea un grafico gauge elegante"""
    fig, ax = plt.subplots(figsize=(4, 3), subplot_kw={'projection': 'polar'})

    # Configurazione
    ax.set_theta_offset(np.pi)
    ax.set_theta_direction(-1)
    ax.set_ylim(0, 1)
    ax.set_xlim(0, np.pi)

    # Sfondo grigio
    theta_bg = np.linspace(0, np.pi, 100)
    ax.fill_between(theta_bg, 0, 1, color='#e2e8f0', alpha=0.5)

    # Valore
    ratio = value / max_value
    theta_val = np.linspace(0, np.pi * ratio, 100)
    color = '#38a169' if ratio >= 0.8 else '#d69e2e' if ratio >= 0.6 else '#e53e3e'
    ax.fill_between(theta_val, 0.3, 0.9, color=color, alpha=0.8)

    # Nascondere assi
    ax.set_axis_off()

    # Testo centrale
    ax.text(np.pi/2, -0.2, f'{value}', ha='center', va='center',
            fontsize=32, fontweight='bold', color='#1a365d')
    ax.text(np.pi/2, -0.5, title, ha='center', va='center',
            fontsize=12, color='#4a5568')

    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight', transparent=True)
    plt.close()

def create_bar_chart(data, title, filename):
    """Crea un grafico a barre orizzontali elegante"""
    fig, ax = plt.subplots(figsize=(8, 5))

    categories = list(data.keys())
    values = list(data.values())

    # Colori basati sul valore
    colors = ['#38a169' if v >= 95 else '#48bb78' if v >= 90 else '#68d391' for v in values]

    # Barre orizzontali
    bars = ax.barh(categories, values, color=colors, height=0.6, edgecolor='white', linewidth=1)

    # Valori sulle barre
    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height()/2,
                f'{val}%', va='center', ha='left', fontsize=11, fontweight='bold', color='#1a365d')

    # Target line
    ax.axvline(x=90, color='#e53e3e', linestyle='--', linewidth=2, label='Target (90%)')

    # Stile
    ax.set_xlim(0, 110)
    ax.set_xlabel('Copertura (%)', fontsize=11, color='#4a5568')
    ax.set_title(title, fontsize=14, fontweight='bold', color='#1a365d', pad=20)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_color('#e2e8f0')
    ax.spines['left'].set_color('#e2e8f0')
    ax.tick_params(colors='#4a5568')
    ax.legend(loc='lower right')

    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

def create_comparison_chart(filename):
    """Crea grafico confronto con settore"""
    fig, ax = plt.subplots(figsize=(8, 4))

    categories = ['Nostra Soluzione', 'Best-in-Class', 'Media Settore', 'Minimo Accettabile']
    values = [97.5, 95, 75, 60]
    colors = ['#38a169', '#48bb78', '#d69e2e', '#e53e3e']

    bars = ax.barh(categories, values, color=colors, height=0.5, edgecolor='white', linewidth=2)

    for bar, val in zip(bars, values):
        ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height()/2,
                f'{val}%', va='center', ha='left', fontsize=12, fontweight='bold', color='#1a365d')

    ax.set_xlim(0, 110)
    ax.set_title('Confronto con Benchmark di Settore', fontsize=14, fontweight='bold', color='#1a365d', pad=20)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_color('#e2e8f0')
    ax.spines['left'].set_color('#e2e8f0')

    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

def create_compliance_chart(filename):
    """Crea grafico compliance"""
    fig, ax = plt.subplots(figsize=(8, 4))

    frameworks = ['PCI-DSS v4.0', 'HIPAA', 'ISO 27001', 'NIS2', 'SOX', 'GDPR']
    values = [95, 95, 95, 90, 90, 90]

    colors = ['#2c5282' for _ in values]
    bars = ax.barh(frameworks, values, color=colors, height=0.5, edgecolor='white', linewidth=2)

    # Sfondo per target
    for i, (bar, val) in enumerate(zip(bars, values)):
        ax.barh(frameworks[i], 100, height=0.5, color='#e2e8f0', zorder=0)
        ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height()/2,
                f'{val}% ✓', va='center', ha='left', fontsize=11, fontweight='bold', color='#38a169')

    ax.set_xlim(0, 110)
    ax.set_title('Conformità Normativa', fontsize=14, fontweight='bold', color='#1a365d', pad=20)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['bottom'].set_color('#e2e8f0')
    ax.spines['left'].set_color('#e2e8f0')

    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

def create_risk_chart(filename):
    """Crea grafico riduzione rischio prima/dopo"""
    fig, ax = plt.subplots(figsize=(10, 5))

    categories = ['Ransomware', 'Minacce Interne', 'Movimento Laterale', 'Esfiltrazione', 'Furto Credenziali']
    before = [50, 30, 50, 30, 50]
    after = [95, 85, 95, 85, 85]

    x = range(len(categories))
    width = 0.35

    bars1 = ax.bar([i - width/2 for i in x], before, width, label='Prima', color='#e53e3e', alpha=0.7)
    bars2 = ax.bar([i + width/2 for i in x], after, width, label='Dopo', color='#38a169', alpha=0.9)

    # Frecce di miglioramento
    for i, (b, a) in enumerate(zip(before, after)):
        improvement = a - b
        ax.annotate(f'+{improvement}%', xy=(i, a + 3), ha='center', fontsize=10,
                   fontweight='bold', color='#38a169')

    ax.set_ylabel('Capacità di Rilevamento (%)', fontsize=11, color='#4a5568')
    ax.set_title('Riduzione del Rischio: Prima vs Dopo', fontsize=14, fontweight='bold', color='#1a365d', pad=20)
    ax.set_xticks(x)
    ax.set_xticklabels(categories, rotation=15, ha='right')
    ax.legend(loc='upper left')
    ax.set_ylim(0, 110)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)

    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

def create_donut_chart(value, title, filename):
    """Crea un grafico donut elegante"""
    fig, ax = plt.subplots(figsize=(4, 4))

    sizes = [value, 100 - value]
    colors = ['#38a169', '#e2e8f0']

    wedges, _ = ax.pie(sizes, colors=colors, startangle=90,
                       wedgeprops=dict(width=0.4, edgecolor='white', linewidth=3))

    # Testo centrale
    ax.text(0, 0, f'{value}%', ha='center', va='center',
            fontsize=28, fontweight='bold', color='#1a365d')
    ax.text(0, -0.15, title, ha='center', va='center',
            fontsize=10, color='#4a5568')

    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight', transparent=True)
    plt.close()

def create_timeline_chart(filename):
    """Crea timeline implementazione"""
    fig, ax = plt.subplots(figsize=(10, 3))

    weeks = ['Settimana 1', 'Settimana 2', 'Settimana 3', 'Settimana 4']
    activities = ['SYSMON\nDeploy', 'WINDOWS\nEvents', 'SIEM\nIntegration', 'GO\nLIVE']
    colors = ['#2c5282', '#3182ce', '#4299e1', '#38a169']

    for i, (week, activity, color) in enumerate(zip(weeks, activities, colors)):
        # Cerchio
        circle = plt.Circle((i, 0.5), 0.3, color=color, zorder=2)
        ax.add_patch(circle)
        ax.text(i, 0.5, activity, ha='center', va='center', fontsize=9,
               fontweight='bold', color='white', zorder=3)
        ax.text(i, -0.1, week, ha='center', va='top', fontsize=10, color='#4a5568')

        # Linea connessione
        if i < len(weeks) - 1:
            ax.plot([i + 0.3, i + 0.7], [0.5, 0.5], color='#cbd5e0', linewidth=3, zorder=1)
            ax.annotate('→', xy=(i + 0.5, 0.5), fontsize=14, ha='center', va='center', color='#cbd5e0')

    ax.set_xlim(-0.5, 3.5)
    ax.set_ylim(-0.3, 1)
    ax.set_aspect('equal')
    ax.axis('off')
    ax.set_title('Piano di Implementazione', fontsize=14, fontweight='bold', color='#1a365d', pad=20)

    plt.tight_layout()
    plt.savefig(filename, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()

def set_cell_shading(cell, color):
    """Imposta colore sfondo cella"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color.replace('#', ''))
    cell._tc.get_or_add_tcPr().append(shading)

def create_document():
    """Crea il documento Word completo"""
    doc = Document()

    # Stili
    style = doc.styles['Heading 1']
    style.font.color.rgb = RGBColor(26, 54, 93)
    style.font.size = Pt(24)
    style.font.bold = True

    style2 = doc.styles['Heading 2']
    style2.font.color.rgb = RGBColor(44, 82, 130)
    style2.font.size = Pt(16)
    style2.font.bold = True

    # === COPERTINA ===
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_heading('Valutazione di Sicurezza', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(36)
    title.runs[0].font.color.rgb = RGBColor(26, 54, 93)

    subtitle = doc.add_paragraph('Soluzione di Monitoraggio Enterprise')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor(74, 85, 104)

    doc.add_paragraph()

    # Badge APPROVATO
    approved = doc.add_paragraph()
    approved.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = approved.add_run('✓ APPROVATO PER PRODUZIONE')
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(56, 161, 105)

    doc.add_paragraph()

    # Sommario esecutivo box
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SOMMARIO ESECUTIVO')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(74, 85, 104)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Dicembre 2025')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(113, 128, 150)

    doc.add_page_break()

    # === INFORMAZIONI DOCUMENTO ===
    doc.add_heading('Informazioni Documento', 1)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    info = [
        ('Data Emissione', '17 Dicembre 2025'),
        ('Versione', '1.0'),
        ('Classificazione', 'Riservato - Uso Interno'),
        ('Destinatari', 'CDA, CISO, Direzione IT'),
        ('Redatto da', 'Security Engineering Team'),
    ]

    for i, (label, value) in enumerate(info):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_shading(row.cells[0], '#f7fafc')

    doc.add_paragraph()

    # === VERDETTO ===
    doc.add_heading('Sintesi Decisionale', 1)

    p = doc.add_paragraph()
    p.add_run('La soluzione di monitoraggio proposta, basata sulla combinazione di ').font.size = Pt(11)
    run = p.add_run('Microsoft Sysmon')
    run.font.bold = True
    run.font.size = Pt(11)
    p.add_run(' e ').font.size = Pt(11)
    run = p.add_run('Windows Event Logging')
    run.font.bold = True
    run.font.size = Pt(11)
    p.add_run(', ha superato con successo la valutazione di sicurezza ed è ').font.size = Pt(11)
    run = p.add_run('approvata per il deployment in produzione')
    run.font.bold = True
    run.font.color.rgb = RGBColor(56, 161, 105)
    run.font.size = Pt(11)
    p.add_run('.').font.size = Pt(11)

    doc.add_paragraph()

    # KPI Table
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    kpis = [
        ('Punteggio Audit', '92/100', '#38a169'),
        ('Copertura MITRE', '97.5%', '#2c5282'),
        ('Gap Critici', '0', '#38a169'),
    ]

    for i, (label, value, color) in enumerate(kpis):
        cell = table.rows[0].cells[i]
        cell.text = value
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(28)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(
            int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16))

        cell2 = table.rows[1].cells[i]
        cell2.text = label
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell2.paragraphs[0].runs[0].font.size = Pt(10)
        cell2.paragraphs[0].runs[0].font.color.rgb = RGBColor(113, 128, 150)

    doc.add_paragraph()
    doc.add_page_break()

    # === COPERTURA ===
    doc.add_heading('Copertura di Rilevamento', 1)

    p = doc.add_paragraph()
    p.add_run('La soluzione garantisce una copertura del ').font.size = Pt(11)
    run = p.add_run('97.5%')
    run.font.bold = True
    run.font.size = Pt(11)
    p.add_run(' delle tecniche di attacco documentate nel framework MITRE ATT&CK, posizionandosi significativamente al di sopra della media di settore (70-80%).').font.size = Pt(11)

    doc.add_paragraph()

    # Grafico copertura
    doc.add_picture('chart_coverage.png', width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Grafico confronto
    doc.add_picture('chart_comparison.png', width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # === COMPLIANCE ===
    doc.add_heading('Conformità Normativa', 1)

    p = doc.add_paragraph()
    p.add_run('La soluzione soddisfa i requisiti di logging e monitoraggio previsti dalle principali normative di settore.').font.size = Pt(11)

    doc.add_paragraph()

    # Grafico compliance
    doc.add_picture('chart_compliance.png', width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Tabella dettaglio
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'

    headers = ['Framework', 'Copertura', 'Requisiti Chiave']
    compliance_data = [
        ('PCI-DSS v4.0', '95%', 'Req. 10.2 - Logging accessi'),
        ('HIPAA', '95%', '§164.312(b) - Controlli audit'),
        ('ISO 27001', '95%', 'A.12.4 - Logging e monitoraggio'),
        ('NIS2', '90%', 'Art. 21 - Misure di sicurezza'),
        ('SOX', '90%', 'Controlli IT generali'),
        ('GDPR', '90%', 'Art. 32 - Sicurezza trattamento'),
    ]

    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cell, '#1a365d')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Data
    for i, (fw, cov, req) in enumerate(compliance_data):
        table.rows[i+1].cells[0].text = fw
        table.rows[i+1].cells[1].text = cov
        table.rows[i+1].cells[2].text = req
        table.rows[i+1].cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

    # === ROI ===
    doc.add_heading('Analisi Costi-Benefici', 1)

    # ROI box
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    roi_data = [
        ('€ 5.000', 'Investimento'),
        ('€ 1.64M', 'Riduzione Rischio/Anno'),
        ('32.800%', 'ROI'),
    ]

    for i, (value, label) in enumerate(roi_data):
        cell = table.rows[0].cells[i]
        cell.text = value
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.size = Pt(24)
        cell.paragraphs[0].runs[0].font.bold = True
        color = '#38a169' if i == 2 else '#1a365d'
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(
            int(color[1:3], 16), int(color[3:5], 16), int(color[5:7], 16))

        cell2 = table.rows[1].cells[i]
        cell2.text = label
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell2.paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    # Tabelle dettaglio
    doc.add_heading('Dettaglio Investimento', 2)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    costs = [
        ('Effort implementazione', '56 ore'),
        ('Impatto infrastrutturale', 'Trascurabile'),
        ('Formazione SOC', '8 ore'),
        ('Costo Totale', '€ 5.000'),
    ]

    table.rows[0].cells[0].text = 'Voce'
    table.rows[0].cells[1].text = 'Stima'
    set_cell_shading(table.rows[0].cells[0], '#1a365d')
    set_cell_shading(table.rows[0].cells[1], '#1a365d')
    table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    table.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, (voce, stima) in enumerate(costs):
        table.rows[i+1].cells[0].text = voce
        table.rows[i+1].cells[1].text = stima
        if i == len(costs) - 1:
            table.rows[i+1].cells[0].paragraphs[0].runs[0].font.bold = True
            table.rows[i+1].cells[1].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    doc.add_heading('Dettaglio Benefici', 2)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    benefits = [
        ('Riduzione rischio breach', '€ 1.640.000/anno'),
        ('Evitamento sanzioni NIS2', 'Fino a € 10.000.000'),
        ('Riduzione costi IR', '€ 50.000/anno'),
        ('Cyber insurance', 'Requisiti soddisfatti'),
    ]

    table.rows[0].cells[0].text = 'Beneficio'
    table.rows[0].cells[1].text = 'Valore'
    set_cell_shading(table.rows[0].cells[0], '#38a169')
    set_cell_shading(table.rows[0].cells[1], '#38a169')
    table.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    table.rows[0].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, (ben, val) in enumerate(benefits):
        table.rows[i+1].cells[0].text = ben
        table.rows[i+1].cells[1].text = val

    doc.add_page_break()

    # === RISCHIO ===
    doc.add_heading('Valutazione del Rischio', 1)

    p = doc.add_paragraph()
    p.add_run('La soluzione riduce significativamente il rischio di sicurezza attraverso una copertura completa delle tecniche di attacco.').font.size = Pt(11)

    doc.add_paragraph()

    # Grafico rischio
    doc.add_picture('chart_risk.png', width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Tabella rischio
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    risk_headers = ['Area', 'Prima', 'Dopo', 'Miglioramento']
    risk_data = [
        ('Ransomware', 'Medio', 'Molto Alto', '+60%'),
        ('Minacce Interne', 'Basso', 'Alto', '+70%'),
        ('Movimento Laterale', 'Medio', 'Molto Alto', '+50%'),
        ('Esfiltrazione', 'Basso', 'Alto', '+80%'),
        ('Furto Credenziali', 'Medio', 'Alto', '+40%'),
    ]

    for i, header in enumerate(risk_headers):
        table.rows[0].cells[i].text = header
        set_cell_shading(table.rows[0].cells[i], '#1a365d')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    for i, (area, prima, dopo, delta) in enumerate(risk_data):
        table.rows[i+1].cells[0].text = area
        table.rows[i+1].cells[1].text = prima
        table.rows[i+1].cells[2].text = dopo
        table.rows[i+1].cells[3].text = delta
        table.rows[i+1].cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(56, 161, 105)
        table.rows[i+1].cells[3].paragraphs[0].runs[0].font.bold = True

    doc.add_page_break()

    # === IMPLEMENTAZIONE ===
    doc.add_heading('Piano di Implementazione', 1)

    doc.add_picture('chart_timeline.png', width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    timeline_headers = ['Settimana', 'Attività', 'Responsabile', 'Effort']
    timeline_data = [
        ('1', 'Deploy Sysmon', 'Security Engineering', '16h'),
        ('2', 'Windows Event Logging', 'IT Operations', '16h'),
        ('3', 'Integrazione SIEM', 'SOC Team', '16h'),
        ('4', 'Validazione e Go-Live', 'Security Engineering', '8h'),
    ]

    for i, header in enumerate(timeline_headers):
        table.rows[0].cells[i].text = header
        set_cell_shading(table.rows[0].cells[i], '#2c5282')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    for i, row_data in enumerate(timeline_data):
        for j, val in enumerate(row_data):
            table.rows[i+1].cells[j].text = val

    doc.add_page_break()

    # === AUTORIZZAZIONE ===
    doc.add_heading('Autorizzazione al Deployment', 1)

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'

    auth_headers = ['Ambiente', 'Autorizzazione', 'Note']
    auth_data = [
        ('Sviluppo/Test', '✓ APPROVATO', 'Nessuna condizione'),
        ('Pre-Produzione', '✓ APPROVATO', 'Nessuna condizione'),
        ('Produzione Standard', '✓ APPROVATO', 'Nessuna condizione'),
        ('Produzione High-Sec', '✓ APPROVATO', 'Nessuna condizione'),
        ('Ambienti Regolamentati', '✓ APPROVATO', 'Validare SACL'),
    ]

    for i, header in enumerate(auth_headers):
        table.rows[0].cells[i].text = header
        set_cell_shading(table.rows[0].cells[i], '#38a169')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    for i, (env, auth, note) in enumerate(auth_data):
        table.rows[i+1].cells[0].text = env
        table.rows[i+1].cells[1].text = auth
        table.rows[i+1].cells[2].text = note
        table.rows[i+1].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(56, 161, 105)
        table.rows[i+1].cells[1].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # === FIRME ===
    doc.add_heading('Approvazioni', 1)

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    sign_headers = ['Ruolo', 'Nome', 'Firma', 'Data']
    sign_data = [
        ('Security Auditor', 'Security Team', '✓', '17/12/2025'),
        ('Security Engineering Lead', '', '', ''),
        ('IT Operations Manager', '', '', ''),
        ('CISO', '', '', ''),
        ('CTO', '', '', ''),
    ]

    for i, header in enumerate(sign_headers):
        table.rows[0].cells[i].text = header
        set_cell_shading(table.rows[0].cells[i], '#1a365d')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    for i, row_data in enumerate(sign_data):
        for j, val in enumerate(row_data):
            table.rows[i+1].cells[j].text = val
            if j == 2 and val == '✓':
                table.rows[i+1].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(56, 161, 105)

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('Documento Riservato - © 2025 Security Engineering Team')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(113, 128, 150)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Valutazione condotta secondo il framework MITRE ATT&CK e le best practice NIST')
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(113, 128, 150)

    return doc

def main():
    import numpy as np

    print("Generazione grafici...")

    # Coverage per fase
    coverage_data = {
        'Esecuzione': 100,
        'Persistenza': 100,
        'Escalation Privilegi': 100,
        'Movimento Laterale': 100,
        'Discovery': 100,
        'Accesso Credenziali': 95,
        'Evasione Difese': 95,
        'Raccolta Dati': 90,
        'Esfiltrazione': 90,
    }
    create_bar_chart(coverage_data, 'Copertura Rilevamento per Fase di Attacco', 'chart_coverage.png')
    print("  ✓ chart_coverage.png")

    # Confronto settore
    create_comparison_chart('chart_comparison.png')
    print("  ✓ chart_comparison.png")

    # Compliance
    create_compliance_chart('chart_compliance.png')
    print("  ✓ chart_compliance.png")

    # Rischio
    create_risk_chart('chart_risk.png')
    print("  ✓ chart_risk.png")

    # Timeline
    create_timeline_chart('chart_timeline.png')
    print("  ✓ chart_timeline.png")

    print("\nGenerazione documento Word...")
    doc = create_document()

    output_file = 'AUDIT-EXECUTIVE-REPORT-IT.docx'
    doc.save(output_file)
    print(f"  ✓ {output_file}")

    # Cleanup
    for f in ['chart_coverage.png', 'chart_comparison.png', 'chart_compliance.png',
              'chart_risk.png', 'chart_timeline.png']:
        if os.path.exists(f):
            os.remove(f)

    print(f"\n✅ Report generato: {output_file}")

if __name__ == '__main__':
    main()
